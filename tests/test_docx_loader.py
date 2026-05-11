"""docx_loader: round-trip a generated .docx through the loader."""

from pathlib import Path

from docx import Document

from pmi_agent.docx_loader import load_docx


def test_load_docx_renders_headings_paragraphs_and_tables(tmp_path: Path) -> None:
    path = tmp_path / "tiny.docx"
    doc = Document()
    doc.add_heading("Section 1", level=1)
    doc.add_paragraph("Hello world.")
    doc.add_heading("Subsection", level=2)
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "2"
    doc.save(path)

    loaded = load_docx(path)

    assert "# Section 1" in loaded.structured_text
    assert "Hello world." in loaded.structured_text
    assert "## Subsection" in loaded.structured_text
    assert "A | B" in loaded.structured_text
    assert "1 | 2" in loaded.structured_text
    assert len(loaded.sha256) == 64
    assert loaded.source_path == path


def test_load_docx_is_deterministic(tmp_path: Path) -> None:
    path = tmp_path / "stable.docx"
    doc = Document()
    doc.add_paragraph("Same content.")
    doc.save(path)

    first = load_docx(path)
    second = load_docx(path)

    assert first.sha256 == second.sha256
    assert first.structured_text == second.structured_text
